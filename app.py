from flask import Flask, render_template, request, redirect, url_for, send_file
import os
import fitz  # PyMuPDF
from gtts import gTTS
import numpy as np
from pdf2image import convert_from_path
import cv2
import matplotlib.pyplot as plt
import keras_ocr
import math
import fitz
import os
import docx
from PIL import Image
import pytesseract
from promptcap import PromptCap
import torch

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def extract_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ''

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()

    return text


def text_to_mp3(text, mp3_path):
    tts = gTTS(text)
    tts.save(mp3_path)


@app.route('/', methods=['GET', 'POST'])
def upload_pdf():
    mp3_path = None

    if request.method == 'POST':
        pdf_file = request.files['pdf_file']
        if pdf_file and pdf_file.filename.endswith('.pdf'):
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_file.filename)
            pdf_file.save(pdf_path)

            ##################
            input_pdf_path = pdf_path
            output_image_folder = "content\output_images"
            output_text_removed_folder = "text_removed_images"
            output_text_extracted_folder = "extracted_text"
            text_removed_folder = "content\text_removed_images"

            original_image_folder = "content\output_images"
            processed_image_folder = "content\text_removed_images"
            output_folder = "final_images"

            os.makedirs(output_image_folder, exist_ok=True)
            os.makedirs(output_text_removed_folder, exist_ok=True)
            os.makedirs(output_text_extracted_folder, exist_ok=True)
            os.makedirs(output_folder, exist_ok=True)

            docx_folder_path = "content\extracted_text"
            image_folder_root = "content"

            context_dict = []
            index_const = 0

            """# PDF to Image"""

            def pdf_to_images(pdf_path, image_folder):
                pdf_document = fitz.open(pdf_path)

                for page_number in range(len(pdf_document)):
                    page = pdf_document[page_number]
                    image = page.get_pixmap()
                    image_path = f"{image_folder}/page_{page_number + 1}.png"
                    image.save(image_path, "PNG")
                pdf_document.close()

            """# Extraction followed by removal of text"""

            def midpoint(x1, y1, x2, y2):
                x_mid = int((x1 + x2) / 2)
                y_mid = int((y1 + y2) / 2)
                return (x_mid, y_mid)

            def inpaint_text(img_path, pipeline, page_number):
                img = keras_ocr.tools.read(img_path)
                prediction_groups = pipeline.recognize([img])
                extracted_text = []
                mask = np.zeros(img.shape[:2], dtype="uint8")
                for box in prediction_groups[0]:
                    x0, y0 = box[1][0]
                    x1, y1 = box[1][1]
                    x2, y2 = box[1][2]
                    x3, y3 = box[1][3]

                    x_mid0, y_mid0 = midpoint(x1, y1, x2, y2)
                    x_mid1, y_mi1 = midpoint(x0, y0, x3, y3)

                    thickness = int(math.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2))

                    cv2.line(mask, (x_mid0, y_mid0), (x_mid1, y_mi1), 255,
                             thickness)
                    inpainted_img = cv2.inpaint(img, mask, 7, cv2.INPAINT_NS)

                    text = box[0]
                    extracted_text.append(text)

                combined_text = ' '.join(extracted_text)
                dox_output_path = f"{output_text_extracted_folder}/page_{page_number + 1}_text.docx"

                doc = docx.Document()
                doc.add_paragraph(combined_text)
                doc.save(dox_output_path)

                return (inpainted_img)

            """# Image extraction"""

            def detect_images_in_page(page, page_number):
                img = cv2.imread(page)
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                edges = cv2.Canny(gray, 50, 150, apertureSize=3)
                contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                min_area = 1000
                detected_images = []

                for contour in contours:
                    if cv2.contourArea(contour) > min_area:
                        x, y, w, h = cv2.boundingRect(contour)
                        detected_images.append(img[y:y + h, x:x + w])

                        output_folder = f"detected_images_page_{page_number + 1}"
                        os.makedirs(output_folder, exist_ok=True)
                        output_path = os.path.join(output_folder, f"image_{i + 1}.jpg")
                        cv2.imwrite(output_path, detected_images)

                return detected_images

            """# Labelling the image"""

            padding = 40

            def process_images(original_image_path, processed_image_path):

                original_image = cv2.imread(original_image_path)
                height, width, _ = original_image.shape
                processed_image = cv2.imread(processed_image_path)
                processed_image = cv2.resize(processed_image, (width, height))
                gray = cv2.cvtColor(original_image, cv2.COLOR_BGR2GRAY)
                _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY)

                edges = cv2.Canny(gray, 50, 150, apertureSize=3)
                contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

                min_area_threshold = 900

                white_fill = np.ones((height, width, 3), dtype=np.uint8) * 255

                for idx, contour in enumerate(contours):
                    area = cv2.contourArea(contour)

                    if area > min_area_threshold:
                        x, y, w, h = cv2.boundingRect(contour)

                        center_x = x + w // 2
                        center_y = y + h // 2

                        text = "context_dict[idx]"
                        # index_const+=1
                        font = cv2.FONT_HERSHEY_SIMPLEX
                        font_scale = 0.5
                        font_color = (0, 0, 0)
                        font_thickness = 1
                        text_size, _ = cv2.getTextSize(text, font, font_scale, font_thickness)

                        text_x = center_x - text_size[0] // 2
                        text_y = center_y + text_size[1] // 2

                        box_x1 = text_x - padding
                        box_y1 = text_y - text_size[1] - padding
                        box_x2 = text_x + text_size[0] + padding
                        box_y2 = text_y + padding
                        cv2.rectangle(processed_image, (box_x1, box_y1), (box_x2, box_y2), (255, 255, 255), -1)

                        cv2.putText(processed_image, text, (text_x, text_y), font, font_scale, font_color,
                                    font_thickness)

                return processed_image

            """# Generating Caption for Image"""

            def generate_text(context, image):
                model = PromptCap(
                    "vqascore/promptcap-coco-vqa")  # also support OFA checkpoints. e.g. "OFA-Sys/ofa-large"

                if torch.cuda.is_available():
                    model.cuda()

                prompt = "please describe this image according to the given context:" + context

                x = model.caption(prompt, image)

                context_dict.append(x)

            def tts(text, language):
                try:
                    tts = gTTS(text, lang=language)
                    tts.save("output.mp3")
                    os.system("start output.mp3")
                except Exception as e:
                    print(f"Error during text-to-speech conversion: {e}")

            def ocr(img):
                filename = img
                img = cv2.imread(filename)
                img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                text = pytesseract.image_to_string(Image.fromarray(img_rgb))
                text = ' '.join(text)
                return text

            def read_aloud(language):
                for page_number in range(len(os.listdir(output_image_folder))):
                    page_image_path = f"final_images/page_{page_number + 1}.png"
                    text_extracted = ocr(page_image_path)
                    tts(text_extracted, language)

            ##################

            ###########**************#################

            if __name__ == "__main__":

                pipeline = keras_ocr.pipeline.Pipeline()

                ### pdf to image ###
                extracted_images = pdf_to_images(input_pdf_path, output_image_folder)

                ### text extraction and removal ###

                for page_number in range(len(os.listdir(output_image_folder))):
                    page_image_path = f"{output_image_folder}/page_{page_number + 1}.png"
                    img_text_removed = inpaint_text(page_image_path, pipeline, page_number)
                    output_path = f"{output_text_removed_folder}/page_{page_number + 1}_text_removed.jpg"
                    cv2.imwrite(output_path, cv2.cvtColor(img_text_removed, cv2.COLOR_BGR2RGB))

                ### image extraction ###
                i = 0
                n = len(sorted(os.listdir(text_removed_folder)))

                for page_number, page_image_path in enumerate(sorted(os.listdir(text_removed_folder))):
                    if i < n - 1:
                        i += 1
                    else:
                        break
                    page_image = f"{text_removed_folder}/page_{page_number + 1}_text_removed.jpg"
                    detected_images = detect_images_in_page(page_image, page_number)

                ### context generation and enhancement ###
                for page_number, doc_path in enumerate(sorted(os.listdir(docx_folder_path))):
                    if doc_path.endswith(".docx"):
                        docx_file_path = os.path.join(docx_folder_path, doc_path)

                        doc = docx.Document(docx_file_path)

                        text = ""
                        for paragraph in doc.paragraphs:
                            text += paragraph.text + "\n"

                        image_folder = os.path.join(image_folder_root, f"detected_images_page_{page_number + 1}")
                        if os.path.exists(image_folder):
                            image_files = [f for f in os.listdir(image_folder) if f.endswith(".jpg")]
                            if image_files:
                                image_path = os.path.join(image_folder, image_files[0])
                                img = cv2.imread(image_path)
                                generate_text(text, img)

                ### image labelling ###

                original_image_files = sorted(os.listdir(original_image_folder))
                processed_image_files = [f for f in os.listdir(text_removed_folder) if f.endswith('.jpg')]

                i = 0
                n = len(sorted(os.listdir(original_image_folder)))
                for page_number, (original_image_file, processed_image_file) in enumerate(
                        zip(original_image_files, processed_image_files)):
                    original_image_path = os.path.join(original_image_folder, original_image_file)
                    processed_image_path = os.path.join(processed_image_folder, processed_image_file)

                    output_image = process_images(processed_image_path, original_image_path)
                    output_image_file_path = f"{output_folder}/page_{page_number + 1}_with_contours.jpg"
                    cv2.imwrite(output_image_file_path, cv2.cvtColor(output_image, cv2.COLOR_BGR2RGB))

            ###########**************#################

            ## image to speech ##
            read_aloud("english")

            # Redirect to the result page
            return redirect(url_for('result'))

    return render_template('upload.html')


@app.route('/result', methods=['GET'])
def result():
    mp3_path = os.path.join(app.config['OUTPUT_FOLDER'], 'output.mp3')
    return render_template('result.html', mp3_path=mp3_path)


@app.route('/download_mp3')
def download_mp3():
    mp3_path = os.path.join(app.config['OUTPUT_FOLDER'], 'output.mp3')
    return send_file(mp3_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)